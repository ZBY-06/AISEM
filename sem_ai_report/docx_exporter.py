# -*- coding: utf-8 -*-
"""Word（python-docx）导出。"""

from __future__ import annotations

import io
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def make_docx_bytes(
    title: str,
    sample_id: str,
    material: str,
    stage: str,
    sinter_temp: str,
    sinter_time: str,
    mag: str,
    scale_info: str,
    caption: str,
    notes: str,
    df: pd.DataFrame,
    full_report_text: str,
    *,
    heating_rate: str = "",
    atmosphere: str = "",
    process_note: str = "",
) -> bytes:
    """生成 Word（.docx）二进制内容；正文已由上层拼接 proxy 与不确定性声明。"""
    doc = Document()
    p_title = doc.add_heading(title, level=0)
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    meta.add_run(f"样品编号：{sample_id}\n")
    meta.add_run(f"材料：{material}\n")
    meta.add_run(f"阶段：{stage}\n")
    meta.add_run(f"烧结温度：{sinter_temp}\n")
    meta.add_run(f"烧结时间：{sinter_time}\n")
    meta.add_run(f"升温速率：{heating_rate or '未提供'}\n")
    meta.add_run(f"气氛：{atmosphere or '未提供'}\n")
    meta.add_run(f"倍率（记录）：{mag}\n")
    meta.add_run(f"比例尺/像素信息：{scale_info or '未填写'}\n")
    meta.add_run(f"工艺备注：{process_note or '无'}\n")
    meta.add_run(f"图注（图片说明）：{caption or '无'}\n")
    meta.add_run(f"备注：{notes or '无'}\n")

    doc.add_heading("图像 proxy 指标表（非严格物理量）", level=1)
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "下表为基于图像的统计 proxy（辅助性草稿），受成像与阈值影响。"
        "dark_area_ratio 仅为暗区面积比例 proxy，禁止称为孔隙率；"
        "edge_density 仅为边缘像素占比 proxy，禁止等同于颗粒边界数量；"
        "sharpness_laplacian_var 仅用于清晰度质控 proxy。"
        "无可靠比例尺或 pixel size 时禁止给出微米单位的粒径、孔径、裂纹长度。"
    )
    disclaimer.runs[0].font.size = Pt(10)

    if df is not None and not df.empty:
        table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for j, col in enumerate(df.columns):
            hdr_cells[j].text = str(col)
        for i, row in enumerate(df.itertuples(index=False), start=1):
            cells = table.rows[i].cells
            for j, val in enumerate(row):
                cells[j].text = str(val)
    else:
        doc.add_paragraph("（无图像或未计算指标）")

    doc.add_heading("报告正文", level=1)
    for line in full_report_text.splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
